"""The docx skill's helpers are executed, not just read.

Two rounds of review found OOXML schema defects in these snippets that every
existing check missed: python-docx accepts out-of-order and duplicated
children without complaint, and the resulting file is still a valid zip, so
"the snippet ran and saved" proves nothing. These tests run the helpers the
skill actually documents and assert the invariants Word enforces.
"""

from __future__ import annotations

import ast
import builtins
import re
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

MEMO_PATH = Path(tempfile.gettempdir()) / "docx_skill_memo_check.docx"

SKILL_MD = (
    Path(__file__).parents[2]
    / "src"
    / "xagent"
    / "skills"
    / "builtin"
    / "docx-report-editorial"
    / "SKILL.md"
)

# w:tcPr / w:tcBorders / w:trPr child order, per ECMA-376 CT_TcPr, CT_TcBorders
# and CT_TrPr. python-docx encodes the same sequences in its _tag_seq tuples.
TCPR_ORDER = (
    "cnfStyle",
    "tcW",
    "gridSpan",
    "hMerge",
    "vMerge",
    "tcBorders",
    "shd",
    "noWrap",
    "tcMar",
    "textDirection",
    "tcFitText",
    "vAlign",
    "hideMark",
)
BORDER_ORDER = (
    "top",
    "left",
    "bottom",
    "right",
    "insideH",
    "insideV",
    "tl2br",
    "tr2bl",
)


def _helpers() -> dict:
    """Execute the skill's own helper snippet and hand back what it defines.

    Reading the code out of SKILL.md is the point: a fix applied to the test
    but not to the documented snippet cannot pass, and agents copy the
    snippet, not this file.
    """
    blocks = re.findall(r"^```python\n(.*?)^```", SKILL_MD.read_text(), re.S | re.M)
    source = next(b for b in blocks if "def shade_cell" in b)
    assert "def set_row_border" in source, "helpers must stay in one block"
    # Empty namespace on purpose — the fence has to carry its own imports.
    namespace: dict = {}
    exec(source, namespace)  # noqa: S102 - executing our own documented snippet
    return namespace


def _child_tags(element) -> list[str]:
    return [child.tag.split("}")[1] for child in element]


def _assert_ordered(element, order: tuple[str, ...], label: str) -> None:
    tags = [t for t in _child_tags(element) if t in order]
    ranks = [order.index(t) for t in tags]
    assert ranks == sorted(ranks), f"{label} out of schema order: {tags}"
    assert len(tags) == len(set(tags)), f"{label} has duplicate children: {tags}"


@pytest.fixture(scope="module")
def helpers() -> dict:
    return _helpers()


def test_shading_then_bordering_one_cell_keeps_tcpr_ordered(helpers) -> None:
    """The order the skill's own sample table uses: zebra-shade the cells,
    then rule the row. w:tcBorders must precede w:shd; appending gave
    ['tcW', 'shd', 'tcBorders'], which Word rejects."""
    table = Document().add_table(rows=3, cols=2)
    cell = table.cell(2, 0)

    helpers["shade_cell"](cell, "E8E5DE")
    helpers["set_row_border"](table.rows[2], "bottom", "0A0A0B")

    tcPr = cell._tc.get_or_add_tcPr()
    _assert_ordered(tcPr, TCPR_ORDER, "w:tcPr")
    tags = _child_tags(tcPr)
    assert tags.index("tcBorders") < tags.index("shd")


def test_bordering_then_shading_one_cell_keeps_tcpr_ordered(helpers) -> None:
    """The reverse order has to hold too -- the helpers must not depend on
    the caller happening to pick the sequence that works."""
    table = Document().add_table(rows=2, cols=2)
    cell = table.cell(0, 0)

    helpers["set_row_border"](table.rows[0], "top", "0A0A0B")
    helpers["shade_cell"](cell, "E8E5DE")

    _assert_ordered(cell._tc.get_or_add_tcPr(), TCPR_ORDER, "w:tcPr")


def test_shading_lands_before_a_later_sibling_already_present(helpers) -> None:
    """shade_cell has to place w:shd correctly on its own, not lean on
    set_row_border's insert to fix the order afterwards. Vertically centring
    a cell first writes w:vAlign, which the schema puts after w:shd."""
    table = Document().add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    helpers["shade_cell"](cell, "0A0A0B")

    tcPr = cell._tc.get_or_add_tcPr()
    _assert_ordered(tcPr, TCPR_ORDER, "w:tcPr")
    tags = _child_tags(tcPr)
    assert tags.index("shd") < tags.index("vAlign")


def test_repeated_calls_replace_rather_than_stack(helpers) -> None:
    """Restyling is idempotent: w:tcPr allows one w:shd and w:tcBorders one
    child per edge, and the last write is what survives."""
    table = Document().add_table(rows=2, cols=2)
    cell = table.cell(1, 0)

    helpers["shade_cell"](cell, "E8E5DE")
    helpers["shade_cell"](cell, "0A0A0B")
    helpers["set_row_border"](table.rows[1], "bottom", "111111")
    helpers["set_row_border"](table.rows[1], "bottom", "FF0000")

    tcPr = cell._tc.get_or_add_tcPr()
    _assert_ordered(tcPr, TCPR_ORDER, "w:tcPr")
    assert tcPr.find(qn("w:shd")).get(qn("w:fill")) == "0A0A0B"

    borders = tcPr.find(qn("w:tcBorders"))
    _assert_ordered(borders, BORDER_ORDER, "w:tcBorders")
    assert borders.find(qn("w:bottom")).get(qn("w:color")) == "FF0000"


def test_every_documented_edge_is_settable(helpers) -> None:
    """Including the diagonals: keying the insert on a tuple that omitted
    tl2br/tr2bl raised ValueError from tuple.index instead of setting them."""
    table = Document().add_table(rows=1, cols=1)

    for edge in BORDER_ORDER:
        helpers["set_row_border"](table.rows[0], edge, "0A0A0B")

    borders = table.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    _assert_ordered(borders, BORDER_ORDER, "w:tcBorders")
    assert set(_child_tags(borders)) == set(BORDER_ORDER)


def test_clearing_a_border_nils_it_in_place(helpers) -> None:
    """Stripping the verticals a full-grid style draws must not disturb the
    horizontal rules already set, nor the element order."""
    table = Document().add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    helpers["set_row_border"](table.rows[0], "bottom", "0A0A0B")
    helpers["clear_cell_border"](cell, "left")
    helpers["clear_cell_border"](cell, "right")

    borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    _assert_ordered(borders, BORDER_ORDER, "w:tcBorders")
    assert borders.find(qn("w:left")).get(qn("w:val")) == "nil"
    assert borders.find(qn("w:bottom")).get(qn("w:color")) == "0A0A0B"


def test_the_validation_snippet_accepts_a_one_page_memo() -> None:
    """The skill advertises memos and letters and forbids padding, so the
    documented check cannot impose a shape minimum: a >10-paragraph or
    >=1-table assertion turns a legitimate short document into an executor
    failure, and the only way to satisfy it is fabricated content."""
    blocks = re.findall(r"^```python\n(.*?)^```", SKILL_MD.read_text(), re.S | re.M)
    source = next(b for b in blocks if "check = Document(" in b)

    memo = Document()
    memo.add_heading("Q3 headcount freeze", level=1)
    memo.add_paragraph("Hiring pauses on 1 October and resumes in January.")
    memo.save(str(MEMO_PATH))

    namespace = {"Document": Document}
    exec(  # noqa: S102 - executing our own documented snippet
        source.replace('"market_expansion_report.docx"', repr(str(MEMO_PATH))),
        namespace,
    )


def test_every_self_contained_fence_imports_what_it_uses() -> None:
    """A fence that opens with its own import line reads as copy-and-run, so
    its import list has to be complete. Listing some of what it uses is the
    worst shape: it looks self-contained and fails at the missing name. Two
    rounds of review found one each (the repeat-header block, then the cover
    block, which imported the two enums it had stopped using and omitted the
    Pt and RGBColor it still called).

    Fences that continue from earlier ones are exempt -- they legitimately
    build on `doc`, `sec`, `table` and `palette` -- so the check is scoped to
    fences that declare imports of their own.
    """
    fences = re.findall(r"^```python\n(.*?)^```", SKILL_MD.read_text(), re.S | re.M)
    carried_over = {"doc", "sec", "table", "row", "cell", "check", "palette"}
    # Helpers an earlier fence defines are carried forward the same way the
    # document objects are -- the skill tells the model to copy them once.
    carried_over |= {
        node.name
        for fence in fences
        for node in ast.walk(ast.parse(fence))
        if isinstance(node, ast.FunctionDef)
    }

    # Static analysis rather than exec: running a fence against stub objects
    # raises AttributeError on the first stub call, which masks every missing
    # name after it. The point is the whole import list, not the first line.
    offenders = []
    for fence in fences:
        if not re.match(r"\s*(from|import)\s", fence):
            continue  # a continuation fence, not advertised as standalone
        tree = ast.parse(fence)
        bound = set(dir(builtins)) | carried_over
        for node in ast.walk(tree):
            if isinstance(node, (ast.Import, ast.ImportFrom)):
                bound |= {a.asname or a.name.split(".")[0] for a in node.names}
            elif isinstance(node, ast.Name) and isinstance(node.ctx, ast.Store):
                bound.add(node.id)
            elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                bound.add(node.name)
                bound |= {a.arg for a in node.args.args}
            elif isinstance(node, ast.ExceptHandler) and node.name:
                bound.add(node.name)
            elif isinstance(node, ast.comprehension):
                bound |= {
                    n.id for n in ast.walk(node.target) if isinstance(n, ast.Name)
                }
        used = {
            n.id
            for n in ast.walk(tree)
            if isinstance(n, ast.Name) and isinstance(n.ctx, ast.Load)
        }
        missing = sorted(used - bound)
        if missing:
            offenders.append((fence.strip().splitlines()[0], missing))

    assert not offenders, f"fences missing imports: {offenders}"


def test_the_prose_does_not_promise_what_the_snippets_omit() -> None:
    """Four rounds of review found prose describing behaviour the code below
    it did not have: a page break where the code opens a section, verticals
    "stripped" by a call that was not there, a cover rule with no rule, an
    eastAsia font pinned against the rule forbidding it. Each was found by
    reading. These pin the ones that are mechanically checkable.
    """
    text = SKILL_MD.read_text()
    fences = "\n".join(re.findall(r"^```python\n(.*?)^```", text, re.S | re.M))

    # The cover: prose names the section break, so the code must open one.
    assert "add_section" in fences
    assert "WD_BREAK.PAGE" not in fences, (
        "a page break next to the section break is the blank-page trap"
    )

    # The table: prose says horizontal rules only, so the verticals the
    # "Table Grid" style draws have to actually be cleared.
    if '"Table Grid"' in fences:
        # A call, not a mention: the helper's own definition contains each edge
        # name, so searching the text passes even with every call deleted.
        sample = next(
            f
            for f in re.findall(r"^```python\n(.*?)^```", text, re.S | re.M)
            if '"Table Grid"' in f
        )
        calls = [
            node
            for node in ast.walk(ast.parse(sample))
            if isinstance(node, ast.Call)
            and isinstance(node.func, ast.Name)
            and node.func.id == "clear_cell_border"
        ]
        assert calls, "the sample never clears the grid's verticals"
        cleared = " ".join(ast.dump(c) for c in calls) + sample
        for vertical in ("left", "right", "insideV"):
            assert vertical in cleared, f"{vertical} border never cleared"

    # Rule 2 puts CJK on platform fallback, so no named face may be pinned.
    assert 'qn("w:eastAsia")' not in fences, "rule 2 forbids pinning a CJK face"

    # The validation snippet must not reimpose a shape floor.
    assert "len(check.paragraphs) > 10" not in fences


def test_the_repeat_header_snippet_stays_single(helpers) -> None:
    """w:trPr permits at most one w:tblHeader; the snippet ran twice -- an
    agent looping over tables or retrying -- used to leave two behind."""
    blocks = re.findall(r"^```python\n(.*?)^```", SKILL_MD.read_text(), re.S | re.M)
    source = next(b for b in blocks if "w:tblHeader" in b)

    # Only `table` is supplied: an agent copying this fence gets whatever the
    # fence itself imports and nothing else. Injecting qn/OxmlElement here is
    # what let a NameError in the published snippet pass review twice.
    document = Document()
    namespace = {"table": document.add_table(rows=2, cols=2)}
    exec(source, namespace)  # noqa: S102 - executing our own documented snippet
    exec(source, namespace)  # noqa: S102 - a retry must not stack a second one

    trPr = namespace["table"].rows[0]._tr.get_or_add_trPr()
    assert len(trPr.findall(qn("w:tblHeader"))) == 1
